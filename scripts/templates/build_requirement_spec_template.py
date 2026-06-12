"""요구사항정의서 워드 양식(backend/templates/requirement_spec.docx) 제작 스크립트.

docxtpl(Jinja) 태그가 삽입된 템플릿을 재현 가능하게 만들기 위한 일회성 제작 도구.
런타임 렌더러는 이 스크립트가 아니라 생성된 .docx 파일을 사용한다.

실행: uv run python scripts/templates/build_requirement_spec_template.py
"""

from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import _Cell
from docx.text.paragraph import Paragraph

OUTPUT_PATH = (
    Path(__file__).resolve().parents[2] / "backend" / "templates" / "requirement_spec.docx"
)

FONT_NAME = "맑은 고딕"
LABEL_SHADE = "D9D9D9"


def set_korean_font(style_or_run, name: str = FONT_NAME) -> None:  # noqa: ANN001
    """한글 폰트는 eastAsia 속성까지 지정해야 워드에서 실제 적용된다."""
    style_or_run.font.name = name
    rpr = style_or_run.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def shade_cell(cell: _Cell, color: str = LABEL_SHADE) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shd)


def style_cell(cell: _Cell, text: str, *, bold: bool = False, center: bool = False) -> None:
    cell.text = text
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0] if para.runs else para.add_run("")
    run.font.size = Pt(10)
    run.font.bold = bold
    set_korean_font(run)


def add_page_number_field(paragraph: Paragraph) -> None:
    """바닥글에 PAGE 필드를 삽입한다 (python-docx 미지원 기능이라 XML 직접 구성)."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def add_tag_paragraph(doc: DocumentType, tag: str) -> None:
    """docxtpl 블록 태그({%p ... %})만 담는 문단을 추가한다 (렌더링 시 문단 자체가 제거됨)."""
    para = doc.add_paragraph(tag)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)


def build_cover(doc: DocumentType) -> None:
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("요구사항정의서")
    run.font.size = Pt(28)
    run.font.bold = True
    set_korean_font(run)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("{{ project_name }}")
    run.font.size = Pt(16)
    set_korean_font(run)

    for _ in range(6):
        doc.add_paragraph()

    info = doc.add_table(rows=4, cols=2)
    info.style = "Table Grid"
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("시스템명", "{{ system_name }}"),
        ("문서번호", "{{ doc_no }}"),
        ("작성자", "{{ author }}"),
        ("작성일", "{{ written_date }}"),
    ]
    for row, (label, value) in zip(info.rows, rows, strict=True):
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(8)
        style_cell(row.cells[0], label, bold=True, center=True)
        shade_cell(row.cells[0])
        style_cell(row.cells[1], value, center=True)
    doc.add_page_break()


def build_revision_table(doc: DocumentType) -> None:
    doc.add_heading("1. 개정 이력", level=1)
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    headers = ["버전", "일자", "작성자", "변경 내용"]
    widths = [Cm(2), Cm(3), Cm(3), Cm(8)]
    for cell, header, width in zip(table.rows[0].cells, headers, widths, strict=True):
        cell.width = width
        style_cell(cell, header, bold=True, center=True)
        shade_cell(cell)
    # docxtpl 행 반복: {%tr%} 태그 행은 렌더링 시 제거된다
    style_cell(table.rows[1].cells[0], "{%tr for rev in revisions %}")
    values = [
        "{{ rev.version }}",
        "{{ rev.revised_date }}",
        "{{ rev.author }}",
        "{{ rev.description }}",
    ]
    for cell, value, width in zip(table.rows[2].cells, values, widths, strict=True):
        cell.width = width
        style_cell(cell, value, center=value != "{{ rev.description }}")
    style_cell(table.rows[3].cells[0], "{%tr endfor %}")


def build_requirement_list(doc: DocumentType) -> None:
    doc.add_heading("2. 요건 목록", level=1)
    table = doc.add_table(rows=4, cols=5)
    table.style = "Table Grid"
    headers = ["요건 ID", "요건명", "구분", "중요도", "비고"]
    widths = [Cm(2.5), Cm(6), Cm(2.5), Cm(2), Cm(3)]
    for cell, header, width in zip(table.rows[0].cells, headers, widths, strict=True):
        cell.width = width
        style_cell(cell, header, bold=True, center=True)
        shade_cell(cell)
    style_cell(table.rows[1].cells[0], "{%tr for req in requirements %}")
    values = [
        "{{ req.req_id }}",
        "{{ req.name }}",
        "{{ req.category }}",
        "{{ req.priority }}",
        "{{ req.note }}",
    ]
    for cell, value, width in zip(table.rows[2].cells, values, widths, strict=True):
        cell.width = width
        style_cell(cell, value, center=value != "{{ req.name }}")
    style_cell(table.rows[3].cells[0], "{%tr endfor %}")


def build_requirement_detail(doc: DocumentType) -> None:
    doc.add_heading("3. 요건 상세", level=1)
    add_tag_paragraph(doc, "{%p for req in requirements %}")
    heading = doc.add_heading("{{ req.req_id }} {{ req.name }}", level=2)
    for run in heading.runs:
        set_korean_font(run)

    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    widths = [Cm(3), Cm(5), Cm(3), Cm(5)]
    grid = [
        ["요건 ID", "{{ req.req_id }}", "구분", "{{ req.category }}"],
        ["요건명", "{{ req.name }}", "중요도", "{{ req.priority }}"],
    ]
    for row, cells in zip(table.rows[:2], grid, strict=True):
        for cell, value, width in zip(row.cells, cells, widths, strict=True):
            cell.width = width
            is_label = value in {"요건 ID", "요건명", "구분", "중요도"}
            style_cell(cell, value, bold=is_label, center=is_label)
            if is_label:
                shade_cell(cell)
    # 설명/비고 행: 라벨 1칸 + 값 3칸 병합
    for row_idx, (label, value) in enumerate(
        [("요건 설명", "{{ req.description }}"), ("비고", "{{ req.note }}")], start=2
    ):
        row = table.rows[row_idx]
        row.cells[0].width = widths[0]
        style_cell(row.cells[0], label, bold=True, center=True)
        shade_cell(row.cells[0])
        merged = row.cells[1].merge(row.cells[2]).merge(row.cells[3])
        style_cell(merged, value)
    doc.add_paragraph()
    add_tag_paragraph(doc, "{%p endfor %}")


def main() -> None:
    doc = Document()

    # 기본/제목 스타일 폰트 통일 및 A4 페이지 설정
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10)
    set_korean_font(normal)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13)]:
        style = doc.styles[style_name]
        style.font.size = Pt(size)
        style.font.color.rgb = None  # 기본 파란색 제거 → 자동(검정)
        set_korean_font(style)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 머리글: 문서번호 + 문서명 / 바닥글: 페이지 번호
    header_para = section.header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_para.add_run("{{ doc_no }} | 요구사항정의서")
    run.font.size = Pt(8)
    set_korean_font(run)

    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dash = footer_para.add_run("- ")
    add_page_number_field(footer_para)
    dash2 = footer_para.add_run(" -")
    for r in (dash, dash2, *footer_para.runs):
        r.font.size = Pt(8)

    build_cover(doc)
    build_revision_table(doc)
    build_requirement_list(doc)
    build_requirement_detail(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"생성 완료: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
