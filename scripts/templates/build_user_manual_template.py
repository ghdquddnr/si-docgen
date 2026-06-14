"""사용자 매뉴얼 워드 양식(backend/templates/user_manual.docx) 제작 스크립트.

docxtpl(Jinja) 태그가 삽입된 템플릿을 재현 가능하게 만드는 일회성 도구.
런타임 렌더러는 이 스크립트가 아니라 생성된 .docx 파일을 사용한다.
섹션→단계 중첩 반복 + 단계별 화면 캡처({{ step.image }}) 자리를 둔다.

실행: uv run python scripts/templates/build_user_manual_template.py
"""

from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import _Cell
from docx.text.paragraph import Paragraph

OUTPUT_PATH = Path(__file__).resolve().parents[2] / "backend" / "templates" / "user_manual.docx"

FONT_NAME = "맑은 고딕"
LABEL_SHADE = "D9D9D9"


def set_korean_font(style_or_run, name: str = FONT_NAME) -> None:  # noqa: ANN001
    style_or_run.font.name = name
    rpr = style_or_run.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def shade_cell(cell: _Cell, color: str = LABEL_SHADE) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shd)


def style_cell(cell: _Cell, text: str, *, bold: bool = False, center: bool = False) -> None:
    cell.text = text
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0] if para.runs else para.add_run("")
    run.font.size = Pt(10)
    run.font.bold = bold
    set_korean_font(run)


def add_page_number_field(paragraph: Paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def add_tag_paragraph(doc: DocumentType, tag: str) -> None:
    """docxtpl 블록 태그({%p ... %})만 담는 문단 (렌더링 시 문단 자체가 제거됨)."""
    para = doc.add_paragraph(tag)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)


def build_cover(doc: DocumentType) -> None:
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("사용자 매뉴얼")
    run.font.size = Pt(28)
    run.font.bold = True
    set_korean_font(run)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("{{ project_name }}")
    run.font.size = Pt(16)
    set_korean_font(run)

    for _ in range(6):
        doc.add_paragraph()

    info = doc.add_table(rows=3, cols=2)
    info.style = "Table Grid"
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("시스템명", "{{ system_name }}"),
        ("작성자", "{{ author }}"),
        ("작성일", "{{ written_date }}"),
    ]
    for row, (label, value) in zip(info.rows, rows, strict=True):
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(8)
        style_cell(row.cells[0], label, bold=True, center=True)
        shade_cell(row.cells[0])
        style_cell(row.cells[1], value, center=True)
    doc.add_page_break()


def build_body(doc: DocumentType) -> None:
    """섹션→단계 중첩 반복. 단계마다 설명/이미지/캡션 문단을 둔다."""
    add_tag_paragraph(doc, "{%p for section in sections %}")
    heading = doc.add_heading("{{ loop.index }}. {{ section.title }}", level=1)
    for run in heading.runs:
        set_korean_font(run)
    desc = doc.add_paragraph("{{ section.description }}")
    set_korean_font(desc.runs[0] if desc.runs else desc.add_run(""))

    add_tag_paragraph(doc, "{%p for step in section.steps %}")
    # 단계 설명 (번호는 단계 루프 index)
    step_para = doc.add_paragraph()
    step_run = step_para.add_run("{{ loop.index }}. {{ step.instruction }}")
    step_run.font.bold = True
    set_korean_font(step_run)
    # 화면 캡처 자리 (이미지 InlineImage 또는 '[화면 캡처]' 플레이스홀더 문자열)
    img_para = doc.add_paragraph("{{ step.image }}")
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_korean_font(img_para.runs[0] if img_para.runs else img_para.add_run(""))
    # 캡션 (있을 때만)
    cap_para = doc.add_paragraph("{% if step.caption %}▲ {{ step.caption }}{% endif %}")
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.runs[0] if cap_para.runs else cap_para.add_run("")
    cap_run.font.size = Pt(9)
    cap_run.font.italic = True
    cap_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    set_korean_font(cap_run)
    add_tag_paragraph(doc, "{%p endfor %}")
    add_tag_paragraph(doc, "{%p endfor %}")


def main() -> None:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.size = Pt(10)
    set_korean_font(normal)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13)]:
        style = doc.styles[style_name]
        style.font.size = Pt(size)
        style.font.color.rgb = None
        set_korean_font(style)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    header_para = section.header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_para.add_run("{{ project_name }} | 사용자 매뉴얼")
    run.font.size = Pt(8)
    set_korean_font(run)

    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dash = footer_para.add_run("- ")
    add_page_number_field(footer_para)
    dash2 = footer_para.add_run(" -")
    for r in (dash, dash2, *footer_para.runs):
        r.font.size = Pt(8)

    build_cover(doc)
    build_body(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"생성 완료: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
