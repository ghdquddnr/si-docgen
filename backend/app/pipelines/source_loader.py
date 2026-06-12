"""원천 문서 파서: docx / pdf / md / txt 에서 텍스트를 추출한다.

LLM 입력으로 쓰일 원천 문서를 형식과 무관하게 SourceDocument 로 통일한다.
"""

from pathlib import Path

from pydantic import BaseModel, Field

from app.exceptions import SourceParseError


class SourceDocument(BaseModel):
    """파싱된 원천 문서 — 형식 무관 공통 표현."""

    filename: str = Field(..., description="원본 파일명")
    text: str = Field(..., description="본문 텍스트 (문단 순서 유지, 빈 문단 제외)")
    tables: list[list[list[str]]] = Field(
        default_factory=list, description="표 목록. 표[행][셀] 구조의 문자열 (docx 만 해당)"
    )


def load_source(path: Path) -> SourceDocument:
    """확장자에 맞는 파서로 원천 문서를 읽는다. 미지원 형식은 SourceParseError."""
    if not path.is_file():
        raise SourceParseError(f"원천 문서가 없습니다: {path}")

    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _load_docx(path)
    if suffix in {".md", ".markdown", ".txt"}:
        return _load_text(path)
    if suffix == ".pdf":
        return _load_pdf(path)
    raise SourceParseError(f"지원하지 않는 원천 문서 형식입니다: {suffix} ({path.name})")


def _load_docx(path: Path) -> SourceDocument:
    from docx import Document

    try:
        document = Document(str(path))
    except Exception as exc:
        raise SourceParseError(f"docx 파일을 열 수 없습니다: {path} ({exc})") from exc

    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    tables = [
        [[cell.text.strip() for cell in row.cells] for row in table.rows]
        for table in document.tables
    ]
    return SourceDocument(filename=path.name, text="\n".join(paragraphs), tables=tables)


def _load_text(path: Path) -> SourceDocument:
    # utf-8-sig: BOM 이 있는 Windows 생성 파일도 허용
    text = path.read_text(encoding="utf-8-sig")
    return SourceDocument(filename=path.name, text=text.strip())


def _load_pdf(path: Path) -> SourceDocument:
    from pypdf import PdfReader

    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        raise SourceParseError(f"pdf 파일을 열 수 없습니다: {path} ({exc})") from exc

    pages = [page.extract_text().strip() for page in reader.pages]
    return SourceDocument(filename=path.name, text="\n".join(p for p in pages if p))
