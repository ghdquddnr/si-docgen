"""사용자 매뉴얼 워드 렌더러 골든 파일 테스트.

고정 픽스처 JSON 으로 파일을 생성한 뒤 python-docx 로 텍스트·이미지를 추출해 검증한다.
이미지가 있는 단계는 InlineImage 로 삽입되고, 없는 참조는 플레이스홀더 문자열로 표시된다.
"""

import json
from pathlib import Path
from typing import Any

import pytest
from docx import Document

from app.renderers.user_manual_renderer import render_user_manual
from app.schemas.user_manual import UserManualDocument

ROOT = Path(__file__).resolve().parents[2]
TEMPLATE_PATH = ROOT / "backend" / "templates" / "user_manual.docx"
FIXTURE_PATH = Path(__file__).parent / "fixtures" / "user_manual_2.json"
SCREENSHOT = Path(__file__).parent / "fixtures" / "manual_screenshot.png"


@pytest.fixture(scope="module")
def fixture_data() -> dict[str, Any]:
    return json.loads(FIXTURE_PATH.read_text(encoding="utf-8"))


@pytest.fixture(scope="module")
def rendered_path(fixture_data: dict[str, Any], tmp_path_factory: pytest.TempPathFactory) -> Path:
    doc = UserManualDocument.model_validate(fixture_data)
    out = tmp_path_factory.mktemp("golden") / "user_manual_out.docx"
    # SCR-001 만 이미지 제공 → SCR-002 는 플레이스홀더로 표시되어야 한다
    return render_user_manual(doc, TEMPLATE_PATH, out, images={"SCR-001": SCREENSHOT})


def _text(path: Path) -> str:
    return "\n".join(p.text for p in Document(str(path)).paragraphs)


def test_표지_프로젝트명(rendered_path: Path, fixture_data: dict[str, Any]) -> None:
    assert fixture_data["project_name"] in _text(rendered_path)


def test_섹션_제목_번호(rendered_path: Path) -> None:
    text = _text(rendered_path)
    assert "1. 로그인" in text
    assert "2. 공지사항 조회" in text


def test_단계_설명_번호(rendered_path: Path) -> None:
    text = _text(rendered_path)
    assert "1. 로그인 화면에서 사용자 ID와 비밀번호를 입력" in text
    assert "2. 비밀번호를 5회 잘못 입력" in text
    # 섹션마다 단계 번호는 1부터 리셋
    assert "1. 메인 화면에서 [공지사항] 메뉴" in text


def test_캡션_표시(rendered_path: Path) -> None:
    text = _text(rendered_path)
    assert "▲ 로그인 화면" in text
    assert "▲ 공지사항 목록" in text


def test_이미지_삽입_및_플레이스홀더(rendered_path: Path) -> None:
    document = Document(str(rendered_path))
    # SCR-001 만 실제 이미지 → 인라인 이미지 1개
    assert len(document.inline_shapes) == 1
    # SCR-002 는 이미지 미제공 → 플레이스홀더 문자열
    text = "\n".join(p.text for p in document.paragraphs)
    assert "[화면 캡처: SCR-002]" in text


def test_이미지_없는_단계는_빈_자리(rendered_path: Path) -> None:
    # 두 번째 단계(screen_ref="")는 이미지도 플레이스홀더도 없어야 한다
    text = _text(rendered_path)
    assert "[화면 캡처: ]" not in text
