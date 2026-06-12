"""요구사항정의서 워드 렌더러 골든 파일 테스트.

고정 픽스처 JSON 으로 파일을 생성한 뒤 python-docx 로 텍스트/표 내용을 추출해
기대값과 비교한다. 바이너리 비교는 사용하지 않는다.
"""

import json
from pathlib import Path
from typing import Any

import pytest
from docx import Document

from app.renderers.docx_renderer import render_requirement_spec
from app.schemas import requirement_spec as rs

ROOT = Path(__file__).resolve().parents[2]
TEMPLATE_PATH = ROOT / "backend" / "templates" / "requirement_spec.docx"
FIXTURE_PATH = Path(__file__).parent / "fixtures" / "requirement_spec_15.json"


@pytest.fixture(scope="module")
def fixture_data() -> dict[str, Any]:
    # Any 사용 사유: json.loads 결과의 원시 dict 를 그대로 다룬다
    return json.loads(FIXTURE_PATH.read_text(encoding="utf-8"))


@pytest.fixture(scope="module")
def rendered(fixture_data: dict[str, Any], tmp_path_factory: pytest.TempPathFactory) -> Document:
    doc = rs.RequirementSpecDocument.model_validate(fixture_data)
    out = tmp_path_factory.mktemp("golden") / "requirement_spec_out.docx"
    return Document(render_requirement_spec(doc, TEMPLATE_PATH, out))


def test_섹션_제목_유지(rendered: Document) -> None:
    h1 = [p.text for p in rendered.paragraphs if p.style.name == "Heading 1"]
    assert h1 == ["1. 개정 이력", "2. 요건 목록", "3. 요건 상세"]


def test_요건_상세_섹션_반복(rendered: Document, fixture_data: dict[str, Any]) -> None:
    headings = [p.text for p in rendered.paragraphs if p.style.name == "Heading 2"]
    expected = [f"{req['req_id']} {req['name']}" for req in fixture_data["requirements"]]
    assert headings == expected


def test_요건_목록_표_행과_요건_ID(rendered: Document, fixture_data: dict[str, Any]) -> None:
    req_list = rendered.tables[2]
    requirements = fixture_data["requirements"]
    assert len(req_list.rows) == 1 + len(requirements)
    ids = [row.cells[0].text for row in req_list.rows[1:]]
    assert ids == [req["req_id"] for req in requirements]
    names = [row.cells[1].text for row in req_list.rows[1:]]
    assert names == [req["name"] for req in requirements]


def test_개정_이력_표_행(rendered: Document, fixture_data: dict[str, Any]) -> None:
    revisions = rendered.tables[1]
    assert len(revisions.rows) == 1 + len(fixture_data["revisions"])
    for row, rev in zip(revisions.rows[1:], fixture_data["revisions"], strict=True):
        assert [c.text for c in row.cells] == [
            rev["version"],
            rev["revised_date"],
            rev["author"],
            rev["description"],
        ]


def test_표지_정보_치환(rendered: Document, fixture_data: dict[str, Any]) -> None:
    cover = rendered.tables[0]
    values = {row.cells[0].text: row.cells[1].text for row in cover.rows}
    assert values == {
        "시스템명": fixture_data["system_name"],
        "문서번호": fixture_data["doc_no"],
        "작성자": fixture_data["author"],
        "작성일": fixture_data["written_date"],
    }
    body_text = "\n".join(p.text for p in rendered.paragraphs)
    assert fixture_data["project_name"] in body_text


def test_요건_상세_표_내용(rendered: Document, fixture_data: dict[str, Any]) -> None:
    detail_tables = rendered.tables[3:]
    requirements = fixture_data["requirements"]
    assert len(detail_tables) == len(requirements)
    for table, req in zip(detail_tables, requirements, strict=True):
        assert table.rows[0].cells[1].text == req["req_id"]
        assert table.rows[1].cells[1].text == req["name"]
        assert table.rows[0].cells[3].text == req["category"]
        assert table.rows[1].cells[3].text == req["priority"]
        assert table.rows[2].cells[1].text == req["description"]


def test_잔여_Jinja_태그_없음(rendered: Document) -> None:
    body = "\n".join(p.text for p in rendered.paragraphs)
    tables = "\n".join(c.text for t in rendered.tables for r in t.rows for c in r.cells)
    assert "{{" not in body and "{%" not in body
    assert "{{" not in tables and "{%" not in tables


def test_머리글_바닥글_유지(rendered: Document, fixture_data: dict[str, Any]) -> None:
    section = rendered.sections[0]
    header_text = "\n".join(p.text for p in section.header.paragraphs)
    assert f"{fixture_data['doc_no']} | 요구사항정의서" in header_text
    footer_xml = section.footer._element.xml
    assert "PAGE" in footer_xml and "fldChar" in footer_xml
