"""웹 사용자 매뉴얼 잡(with_user_manual) e2e 테스트 (LLM 모킹)."""

import io
import json
from collections.abc import Iterator
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.api.main import app
from app.config import get_settings
from app.db.base import Base
from app.db.models import Job, JobStatus
from app.db.session import SessionLocal, rebind_engine

ROOT = Path(__file__).resolve().parents[2]
SCREENSHOT = ROOT / "tests" / "golden" / "fixtures" / "manual_screenshot.png"

SCENARIO: dict[str, Any] = {
    "project_name": "P",
    "system_name": "S",
    "author": "A",
    "written_date": "2026-06-15",
    "unit_test_cases": [
        {
            "tc_id": "TC-001",
            "req_id": "REQ-001",
            "category_major": "공통",
            "category_minor": "로그인",
            "scenario_name": "정상",
            "test_steps": ["단계"],
            "expected_result": "결과",
        }
    ],
    "integration_test_cases": [],
}

USER_MANUAL: dict[str, Any] = {
    "project_name": "P",
    "system_name": "S",
    "author": "A",
    "written_date": "2026-06-15",
    "sections": [
        {
            "title": "로그인",
            "description": "시스템 로그인 방법",
            "steps": [
                {
                    "instruction": "ID/PW 입력 후 로그인",
                    "screen_ref": "SCR-001",
                    "caption": "로그인 화면",
                },
                {"instruction": "메인으로 이동", "screen_ref": "", "caption": ""},
            ],
        }
    ],
}


@pytest.fixture
def client(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Iterator[TestClient]:
    monkeypatch.setenv("SIDOCGEN_STORAGE_DIR", str(tmp_path / "jobs"))
    get_settings.cache_clear()
    engine = rebind_engine(f"sqlite:///{tmp_path / 't.db'}")
    Base.metadata.create_all(engine)

    def fake(prompt: str, *, system: str | None = None, json_schema=None, model=None) -> str:
        payload = USER_MANUAL if system and "테크니컬 라이터" in system else SCENARIO
        return json.dumps(payload, ensure_ascii=False)

    monkeypatch.setattr("app.llm.generate.complete_json", fake)
    yield TestClient(app)
    get_settings.cache_clear()


def _manual_job(client: TestClient) -> str:
    resp = client.post(
        "/jobs",
        files={"file": ("req.md", b"REQ-001", "text/markdown")},
        data={"with_user_manual": "true"},
    )
    assert resp.status_code == 201
    assert resp.json()["with_user_manual"] is True
    return resp.json()["id"]


def test_매뉴얼_잡_저장(client: TestClient) -> None:
    job_id = _manual_job(client)
    assert client.get(f"/jobs/{job_id}").json()["status"] == "succeeded"
    with SessionLocal() as db:
        job = db.get(Job, job_id)
        assert job.status is JobStatus.SUCCEEDED
        assert job.user_manual_json is not None
        assert job.user_manual_json["sections"][0]["title"] == "로그인"


def test_매뉴얼_조회(client: TestClient) -> None:
    job_id = _manual_job(client)
    resp = client.get(f"/jobs/{job_id}/user-manual")
    assert resp.status_code == 200
    assert resp.json()["sections"][0]["steps"][0]["screen_ref"] == "SCR-001"


def test_매뉴얼_편집_저장(client: TestClient) -> None:
    job_id = _manual_job(client)
    edited = dict(USER_MANUAL)
    edited["sections"] = [
        {
            "title": "수정됨",
            "description": "",
            "steps": [{"instruction": "수정 단계", "screen_ref": "", "caption": ""}],
        }
    ]
    resp = client.put(f"/jobs/{job_id}/user-manual", json=edited)
    assert resp.status_code == 200
    assert client.get(f"/jobs/{job_id}/user-manual").json()["sections"][0]["title"] == "수정됨"


def test_매뉴얼_편집_스키마위반_422(client: TestClient) -> None:
    job_id = _manual_job(client)
    bad = dict(USER_MANUAL)
    bad["sections"] = []  # 섹션 0건 (min_length=1 위반)
    assert client.put(f"/jobs/{job_id}/user-manual", json=bad).status_code == 422


def test_화면캡처_업로드_여부_조회(client: TestClient) -> None:
    job_id = _manual_job(client)
    refs = client.get(f"/jobs/{job_id}/manual-images")
    assert refs.status_code == 200
    assert refs.json() == {"SCR-001": False}  # screen_ref 있는 단계만, 아직 미업로드


def test_없는_screen_ref_업로드_404(client: TestClient) -> None:
    job_id = _manual_job(client)
    resp = client.post(
        f"/jobs/{job_id}/manual-images/SCR-999",
        files={"file": ("a.png", SCREENSHOT.read_bytes(), "image/png")},
    )
    assert resp.status_code == 404


def test_지원하지않는_이미지형식_400(client: TestClient) -> None:
    job_id = _manual_job(client)
    resp = client.post(
        f"/jobs/{job_id}/manual-images/SCR-001",
        files={"file": ("a.txt", b"not an image", "text/plain")},
    )
    assert resp.status_code == 400


def test_캡처_업로드_후_렌더_docx_다운로드(client: TestClient) -> None:
    job_id = _manual_job(client)
    up = client.post(
        f"/jobs/{job_id}/manual-images/SCR-001",
        files={"file": ("shot.png", SCREENSHOT.read_bytes(), "image/png")},
    )
    assert up.status_code == 201
    assert client.get(f"/jobs/{job_id}/manual-images").json()["SCR-001"] is True

    render = client.post(f"/jobs/{job_id}/render")
    assert render.status_code == 200
    assert "user_manual" in render.json()["downloads"]

    dl = client.get(f"/jobs/{job_id}/download/user_manual")
    assert dl.status_code == 200
    doc = Document(io.BytesIO(dl.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "로그인" in text
    # 업로드한 캡처가 삽입되어 문서에 이미지가 1개 이상 존재한다
    assert len(doc.inline_shapes) >= 1


def test_비매뉴얼_잡은_409(client: TestClient) -> None:
    resp = client.post("/jobs", files={"file": ("req.md", b"REQ-001", "text/markdown")})
    job_id = resp.json()["id"]
    assert resp.json()["with_user_manual"] is False
    assert client.get(f"/jobs/{job_id}/user-manual").status_code == 409


def test_목업_이미지_자동_추출_및_렌더(client: TestClient) -> None:
    job_id = _manual_job(client)
    # 수동 캡처 업로드 없이 render 시 use_mockup_images=true를 전달
    # 이 경우 Pillow Fallback에 의해 SCR-001.png 더미가 생성되어 주입됨
    render = client.post(f"/jobs/{job_id}/render?use_mockup_images=true")
    assert render.status_code == 200
    assert "user_manual" in render.json()["downloads"]
    assert "user-manual-pdf" in render.json()["downloads"]

    dl = client.get(f"/jobs/{job_id}/download/user_manual")
    assert dl.status_code == 200
    doc = Document(io.BytesIO(dl.content))
    # 더미 목업 이미지가 정상 주입되었으므로 inline_shapes 개수가 1개 이상임
    assert len(doc.inline_shapes) >= 1


def test_매뉴얼_PDF_다운로드(client: TestClient) -> None:
    job_id = _manual_job(client)
    # 렌더링을 먼저 수행해야 다운로드 가능
    render = client.post(f"/jobs/{job_id}/render")
    assert render.status_code == 200
    assert "user-manual-pdf" in render.json()["downloads"]

    # 테스트 환경이므로 soffice가 없더라도 더미 PDF 가 정상 리턴됨
    dl = client.get(f"/jobs/{job_id}/download/user-manual-pdf")
    assert dl.status_code == 200
    assert dl.headers["content-type"] == "application/pdf"
    assert b"%PDF-1.4" in dl.content
