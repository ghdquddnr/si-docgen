"""사용자 매뉴얼 LLM 생성 파이프라인 테스트 (LLM 모킹)."""

import json
import shutil
from pathlib import Path
from typing import Any

import pytest
from docx import Document

from app import cli
from app.pipelines.generate_user_manual import (
    collect_images,
    generate_and_render_user_manual,
    generate_user_manual,
)

ROOT = Path(__file__).resolve().parents[2]
INPUT = ROOT / "tests" / "fixtures" / "sources" / "sample_requirements.md"
SCREENSHOT = ROOT / "tests" / "golden" / "fixtures" / "manual_screenshot.png"

MOCK_MANUAL: dict[str, Any] = {
    "project_name": "P",
    "system_name": "S",
    "author": "A",
    "written_date": "2026-06-14",
    "sections": [
        {
            "title": "로그인",
            "description": "로그인 방법",
            "steps": [
                {
                    "instruction": "ID/PW 를 입력하고 로그인 버튼을 클릭합니다.",
                    "screen_ref": "SCR-001",
                    "caption": "로그인 화면",
                },
                {"instruction": "오류 시 메시지를 확인합니다.", "screen_ref": "", "caption": ""},
            ],
        },
        {
            "title": "공지사항",
            "description": "공지 확인",
            "steps": [
                {
                    "instruction": "공지사항 메뉴를 클릭합니다.",
                    "screen_ref": "SCR-002",
                    "caption": "공지 목록",
                }
            ],
        },
    ],
}

COVER = {"project_name": "P", "system_name": "S", "author": "A", "written_date": "2026-06-14"}


@pytest.fixture
def mock_llm(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(
        "app.llm.generate.complete_json",
        lambda *a, **k: json.dumps(MOCK_MANUAL, ensure_ascii=False),
    )


def test_사용자매뉴얼_생성_및_검증(mock_llm: None) -> None:
    doc = generate_user_manual(INPUT, **COVER)
    assert len(doc.sections) == 2
    assert doc.sections[0].steps[0].screen_ref == "SCR-001"


def test_진행_콜백_단계_통지(mock_llm: None) -> None:
    stages: list[str] = []
    generate_user_manual(INPUT, **COVER, on_progress=stages.append)
    assert stages == ["parsing", "generating"]


def test_generate_and_render_플레이스홀더(mock_llm: None, tmp_path: Path) -> None:
    # 캡처 이미지 미제공 → 모든 화면 자리는 플레이스홀더, 인라인 이미지 0개
    result = generate_and_render_user_manual(INPUT, tmp_path, **COVER)
    assert result.user_manual_path.is_file()
    assert result.section_count == 2
    assert result.step_count == 3
    document = Document(str(result.user_manual_path))
    assert len(document.inline_shapes) == 0
    text = "\n".join(p.text for p in document.paragraphs)
    assert "[화면 캡처: SCR-001]" in text


def test_collect_images_파일명_매칭(tmp_path: Path) -> None:
    shutil.copy(SCREENSHOT, tmp_path / "SCR-001.png")
    shutil.copy(SCREENSHOT, tmp_path / "SCR-003.jpg")
    found = collect_images(tmp_path, ["SCR-001", "SCR-002", "SCR-003", ""])
    assert set(found) == {"SCR-001", "SCR-003"}  # SCR-002 파일 없음, "" 무시
    assert found["SCR-001"].name == "SCR-001.png"


def test_images_dir_로_캡처_삽입(mock_llm: None, tmp_path: Path) -> None:
    images_dir = tmp_path / "shots"
    images_dir.mkdir()
    shutil.copy(SCREENSHOT, images_dir / "SCR-001.png")  # SCR-001 만 제공
    out = tmp_path / "out"
    result = generate_and_render_user_manual(INPUT, out, **COVER, images_dir=images_dir)
    assert result.image_count == 1
    document = Document(str(result.user_manual_path))
    assert len(document.inline_shapes) == 1  # SCR-001 삽입
    text = "\n".join(p.text for p in document.paragraphs)
    assert "[화면 캡처: SCR-002]" in text  # SCR-002 는 플레이스홀더 유지


def test_cli_user_manual_종료코드_0(mock_llm: None, tmp_path: Path) -> None:
    code = cli.main(["user-manual", "--input", str(INPUT), "--output", str(tmp_path)])
    assert code == 0
    assert (tmp_path / "user_manual.docx").is_file()


def test_cli_user_manual_images_dir(mock_llm: None, tmp_path: Path) -> None:
    images_dir = tmp_path / "shots"
    images_dir.mkdir()
    shutil.copy(SCREENSHOT, images_dir / "SCR-001.png")
    out = tmp_path / "out"
    code = cli.main(
        [
            "user-manual",
            "--input",
            str(INPUT),
            "--output",
            str(out),
            "--images-dir",
            str(images_dir),
        ]
    )
    assert code == 0
    assert len(Document(str(out / "user_manual.docx")).inline_shapes) == 1
